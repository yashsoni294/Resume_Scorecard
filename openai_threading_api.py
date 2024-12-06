from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse, FileResponse
from PyPDF2 import PdfReader
from docx import Document
import win32com.client
import os
import openai
import io
import zipfile
import pandas as pd
import time
import tkinter as tk
from tkinter import filedialog
from dotenv import load_dotenv
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from datetime import datetime
import threading
from queue import Queue

# Load API Key
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")

openai.api_key = API_KEY


TEMPLATES = {
    "job_description" : """
        
        The text below is a job description:
        {job_description_text}

        Your task is to analyze the job description and extract critical aspects to evaluate a candidate's suitability effectively. Organize the extracted information into structured categories as outlined below. Ensure conciseness and avoid including assumptions or unnecessary details. The structured output will form the foundation for precise scoring.

        1. Candidate Profile
            1.1 Job-Related Keywords:

            Extract highly relevant keywords and phrases, focusing on essential skills, tools, technologies, and qualifications.
            Highlight terms that frequently appear, emphasizing the primary focus areas for the role.
            1.2 Relevant Past Roles and Responsibilities:

            Identify specific roles (e.g., "Project Manager," "Data Analyst") and responsibilities directly relevant to the role described.
            Highlight areas where past experience aligns with the position key objectives.
            1.3 Actionable Responsibilities:

            List clear, measurable, and action-oriented expectations (e.g., "Design and implement X system," "Lead Y project team") to define success in this role.
        2. Experience Requirements
            2.1 Years of Experience:

            Specify the required and preferred experience levels, distinguishing between mandatory and desirable years in relevant fields.
            2.2 Technical Skills:

            Provide a categorized list of required and preferred technical skills, specifying domain-specific tools, programming languages, platforms, or methodologies.
            Highlight core skills critical for the role versus those that are supplementary.
            2.3 Soft Skills and Interpersonal Abilities:

            List required soft skills (e.g., leadership, problem-solving) and interpersonal abilities (e.g., teamwork, collaboration).
            Include any role-specific examples mentioned, such as "strong stakeholder communication skills."

        3. Educational Qualifications and Certifications
            3.1 Minimum Educational Qualifications:

            State the explicit educational requirements for the role (e.g., "Bachelors degree in Computer Science").
            Differentiate between mandatory and preferred qualifications.
            3.2 Certifications and Specialized Training:

            Highlight certifications, licenses, or training programs required or preferred (e.g., "PMP certification," "AWS Certified Solutions Architect").
            Include both general and domain-specific certifications if applicable.

        Output Format:
        Organize the extracted information in bullet-point format under the categories listed above. Ensure the content is:

        Directly aligned with the job description.
        Actionable and structured to facilitate accurate scoring.
        Free from redundant details or assumptions.
        """ ,

    "resume" : """
        The text below is a resume:
        {resume_text}

        Your task is to extract critical information from the resume, focusing only on its content without making assumptions or adding external details. The extracted details should be structured, concise, and actionable to support effective scoring. Use the categories below for organization:

        1. Candidate Profile
            1.1 Keywords Identified:

            Extract relevant keywords that reflect the candidate's skills, roles, expertise, and domain knowledge.
            Highlight recurring themes or terms indicative of specialization or focus areas.
            1.2 Summary of Past Roles:

            Summarize the candidate primary roles, emphasizing key responsibilities and measurable achievements.
            Include details about progression or diversity in roles where mentioned (e.g., growth from Analyst to Manager).
            1.3 Measurable Achievements:

            Identify specific, quantifiable accomplishments (e.g., "Increased revenue by X%," "Reduced costs by Y%").
            Highlight the use of action-oriented language (e.g., "Led," "Implemented," "Designed").
        2. Experience Details

            2.1 Total Years of Experience:

            Indicate the total years of professional experience and the industries or domains the candidate has worked in.
            Include any explicit references to seniority (e.g., "5+ years in project management").
            2.2 Technical Skills and Proficiencies:

            Extract technical skills explicitly mentioned (e.g., tools, programming languages, platforms) and categorize them as core or supplementary.
            Include details of certifications or work examples that validate these skills.
            2.3 Soft Skills and Team Contributions:

            Highlight references to soft skills (e.g., problem-solving, adaptability) and team-related contributions (e.g., collaboration, mentoring).
            Focus on examples that demonstrate these abilities, such as leadership roles or cross-functional projects.
        3. Educational Qualifications and Certifications

            3.1 Educational Background:

            Note the highest qualification achieved, field of study, and any notable academic honors or achievements.
            Include additional qualifications that may complement the role.
            3.2 Certifications and Professional Training:

            List certifications, training programs, and licenses, specifying their relevance to the candidate's field or the role in question.
            Highlight certifications that indicate advanced expertise or specialization (e.g., "AWS Certified Solutions Architect").

        Output Format:
        Present the extracted details in the following format:

        Category: Subcategory/Point (e.g., Candidate Profile: Measurable Achievements).
        Use bullet points or short, clear sentences for each item.
        Ensure alignment with the resume content without adding interpretations or assumptions.

        """ , 
    "score" : """
        Your task is to evaluate the alignment between the provided resume and job description by analyzing three critical sections: Candidate Profile, Experience, and Educational Qualifications and Certifications. Based on your evaluation, assign a final score between 0 and 100, reflecting the overall suitability of the candidate for the job.

        Inputs:
        Resume Text:
        {resume_text}

        Job Description Text:
        {job_description}

        Scoring Guidelines:
        Evaluate the resume against the job description using the criteria outlined below. Assign marks in each category, calculate the total, and round the final score to the nearest whole number.

        1. Candidate Profile (Max 15 Marks)
            1.1 Job-Related Keywords (Max 5 Marks):
                5 Points: Resume includes all highly relevant keywords, indicating strong alignment with job requirements.
                3 Points: Resume includes many relevant keywords but misses some critical ones.
                1 Points: Resume includes few relevant keywords or misses key terms.
            1.2 Relevance of Past Roles to Job Description (Max 5 Marks):
                5 Points: Past roles and responsibilities strongly align with the job description.
                3 Points: Moderate alignment, with partial overlap in roles and responsibilities.
                1 Points: Limited relevance or weak alignment.
            1.3 Clarity of Responsibilities (Max 5 Marks):
                5 Points: Responsibilities are clearly defined using action words (e.g., "Developed," "Managed") with measurable outcomes.
                3 Points: Responsibilities are described but lack clear action words or measurable outcomes.
                1 Points: Responsibilities are vague or generic.

        2. Experience Section (Max 65 Marks)
            2.1 Years of Experience (Max 15 Marks):
                15 Points: Meets or exceeds the required years of experience.
                10 Points: Slightly below the required years but with relevant experience.
                5 Points: Limited relevance or inadequate years of experience.
            2.2 Matching Technical Skills (Max 40 Marks):
                40 Points: All technical skills mentioned in the job description are evident, supported by examples or certifications.
                25 Points: Most technical skills are evident, but examples or certifications are missing.
                15 Points: Some technical skills align, but several are missing.
                5 Points: Minimal or no alignment with the required technical skills.
            2.3 Communication and Teamwork (Max 10 Marks):
                10 Points: Strong evidence of soft skills, supported by examples (e.g., "Led a team of 5," "Facilitated cross-department collaboration").
                7 Points: Mentions soft skills but lacks specific examples.
                3 Points: Minimal or generic mention of soft skills.

        3. Educational Qualifications and Certifications (Max 20 Marks)
            3.1 Minimum Educational Qualifications (Max 15 Marks):
                15 Points: Meets or exceeds the educational qualifications specified in the job description.
                10 Points: Meets basic qualifications but lacks advanced or preferred qualifications.
                5 Points: Does not fully meet the educational qualifications.
            3.2 Additional Certifications/Training Programs (Max 5 Marks):
                5 Points: Certifications/training are directly relevant to the job description (e.g., industry-specific certifications).
                3 Points: Certifications or training are partially relevant to the job description.
                1 Point: No additional certifications or irrelevant certifications.

        Additional Refinements:
            Ensure that scoring accounts for both the breadth and depth of alignment between the resume and job description.
            Emphasize evidence-backed qualifications and experience to avoid scoring inflated or unsupported claims.
        Output:
            Provide the final calculated score as a single whole number (0 – 100) with no additional explanation or text. If you are not able to score the resume then you can give 0 score to the resume.
        """

}

app = FastAPI()


def get_conversation_openai(template, model="gpt-4o-mini", temperature=0.1, max_tokens=None):

    """
    Creates a function that interacts with the OpenAI model based on a provided template.
    
    Args:
        template (str): A string template for generating prompts dynamically. The template should use placeholders
                        that can be filled with dynamic input values.
        model (str, optional): The name of the OpenAI model to use for generating responses. Defaults to "gpt-4o-mini".
        temperature (float, optional): Sampling temperature to control the randomness of the response. 
                                       Lower values make output more focused and deterministic. Defaults to 0.1.
        max_tokens (int, optional): The maximum number of tokens to include in the response. Defaults to None, 
                                    allowing the model to determine the length.

    Returns:
        function: A callable function that takes a dictionary of inputs, formats the prompt based on the template,
                  and interacts with the OpenAI model to generate a response.

    """
 
    # Define a nested function to handle API interaction
    def call_openai_model(inputs):
        """
        Invokes the OpenAI model using the provided template and inputs.
        
        Args:
            inputs (dict): A dictionary containing values for the placeholders in the template.

        Returns:
            str: The content of the response generated by the OpenAI model.
        """
        # Generate the prompt by formatting the template with the provided inputs
        prompt = PromptTemplate.from_template(template).format(**inputs)
        # Call the OpenAI Chat API to generate a response
        response = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "system", "content": prompt}],
            temperature=temperature,
            max_tokens=max_tokens
        )
        # Extract and return the content of the response
        return response["choices"][0]["message"]["content"]
    
    # Return the nested function for reuse
    return call_openai_model

def read_pdf(file: io.BytesIO):
    """Extract text from a PDF file."""
    pdf_reader = PdfReader(file)
    extracted_text = ""
    for page in pdf_reader.pages:
        extracted_text += page.extract_text()
    return extracted_text.strip()

def read_docx(file: io.BytesIO):
    """Extract text from a DOCX file."""
    document = Document(file)
    extracted_text = ""
    for paragraph in document.paragraphs:
        extracted_text += paragraph.text + "\n"
    return extracted_text.strip()

def read_doc(file_path: str):
    """
    Extract text from a DOC file using COM automation (Windows only).
    """
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(file_path))
        extracted_text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        return extracted_text.strip()
    except Exception as e:
        return f"Error processing DOC file: {str(e)}"

def read_txt(file: io.BytesIO):
    """Extract text from a plain text file."""
    try:
        contents = file.read()
        return contents.decode("utf-8").strip()
    except Exception as e:
        return f"Error processing TXT file: {str(e)}"

def threaded_resume_processor(queue, job_description, results):
    conversation_resume = get_conversation_openai(TEMPLATES["resume"])
    conversation_score = get_conversation_openai(TEMPLATES["score"])
    while not queue.empty():
        filename, data = queue.get()
        resume_response = conversation_resume({"resume_text": data["content"]})
        results[filename] = {
            "resume_key_aspect": resume_response,
            "resume_score": conversation_score({
                "resume_text": resume_response,
                "job_description": job_description
            })
        }


@app.post("/upload-files/")
async def upload_files(job_description: str, files: list[UploadFile] = File(...)):
    response_data = {}

    conversation_jd = get_conversation_openai(TEMPLATES["job_description"])
    jd_response = conversation_jd({"job_description_text": job_description})
    processed_jd = jd_response
    print(processed_jd)

    # Create extracted_files directory if it doesn't exist
    extract_path = "extracted_files"
    os.makedirs(extract_path, exist_ok=True)

    for file in files:
        try:
            # Fallback to file extension if MIME type is not reliable
            file_extension = file.filename.split(".")[-1].lower()

            # Generate a unique file name using a timestamp
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
            file_name = file.filename
            unique_filename = f"{timestamp}_{file_name}"
            file_path = os.path.join(extract_path, unique_filename)

            if file.content_type == "application/zip" or file_extension == "zip":
                zip_data = await file.read()
                zip_file = io.BytesIO(zip_data)
                extracted_files = []
                with zipfile.ZipFile(zip_file, 'r') as z:
                    # Extract files
                    file_name_list = z.namelist()
                    for original_file_name in file_name_list:
                        # Generate a unique file name for each file in the ZIP
                        timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
                        unique_file_name = f"{timestamp}_{original_file_name}"
                        z.extract(original_file_name, extract_path)
                        # Rename the extracted file to its unique name
                        os.rename(os.path.join(extract_path, original_file_name),
                                  os.path.join(extract_path, unique_file_name))
                        extracted_files.append(unique_file_name)

                pdf_contents = {}
                i = 0
                for file_name in extracted_files:
                    original_name = file_name_list[i]
                    i += 1
                    file_path = os.path.join(extract_path, file_name)
                    if file_name.endswith(".pdf"):
                        file_path = os.path.join(extract_path, file_name)
                        with open(file_path, "rb") as pdf_file:
                            reader = PdfReader(pdf_file)
                            text = ""
                            for page in reader.pages:
                                text += page.extract_text()
                            response_data[original_name] = {"content": text, "file_path": file_path}
                    
                    # Process TXT files
                    elif file_name.endswith(".txt"):
                        with open(file_path, "r", encoding="utf-8") as txt_file:
                            text = txt_file.read()
                            response_data[original_name] = {"content": text, "file_path": file_path}

                    elif file_name.endswith(".docx"):
                        try:
                            doc = Document(file_path)
                            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            response_data[original_name] = {"content": text, "file_path": file_path}
                        except Exception as e:
                            response_data[original_name] = {"content": e, "file_path": file_path}

                    # Process DOC files
                    elif file_name.endswith(".doc"):
                        try:
                            word = win32com.client.Dispatch("Word.Application")
                            doc = word.Documents.Open(os.path.abspath(file_path))
                            text = doc.Content.Text
                            doc.Close()
                            word.Quit()
                            response_data[original_name] = {"content": text, "file_path": file_path}
                        except Exception as e:
                            response_data[original_name] = {"content": f"Error reading DOC file: {str(e)}", "file_path": file_path}

            else:
                # Save individual file to extracted_files directory
                file_content = await file.read()
                
                with open(file_path, "wb") as f:
                    f.write(file_content)

                # Process based on file type
                if file_extension == "pdf":
                    with open(file_path, "rb") as pdf_file:
                        reader = PdfReader(pdf_file)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text()
                        response_data[file_name] = {"content": text}
                
                elif file_extension == "txt":
                    with open(file_path, "r", encoding="utf-8") as txt_file:
                        text = txt_file.read()
                        response_data[file_name] = {"content": text}
                
                elif file_extension == "docx":
                    try:
                        doc = Document(file_path)
                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                        response_data[file_name] = {"content": text}
                    except Exception as e:
                        response_data[file_name] = {"content": str(e)}
                
                elif file_extension == "doc":
                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        doc = word.Documents.Open(os.path.abspath(file_path))
                        text = doc.Content.Text
                        doc.Close()
                        word.Quit()
                        response_data[file_name] = {"content": text}
                    except Exception as e:
                        response_data[file_name] = {"content": f"Error reading DOC file: {str(e)}"}
                
                else:
                    response_data[file_name] = {
                        "error": "Unsupported file type. Supported formats: .pdf, .docx, .doc, .txt, .zip, .rar"
                    }
                # Add file path to the response data
                response_data[file_name]["file_path"] = file_path

        except Exception as e:
            response_data[file.filename] = {
                "error": str(e)
            }
            
    resume_queue = Queue()
    results = {}

    # Populate the queue with resumes
    for filename, data in response_data.items():
        resume_queue.put((filename, data))

    # Create and start threads
    threads = []
    for _ in range(3):  # Limit to 5 threads or number of resumes
        thread = threading.Thread(target=threaded_resume_processor, args=(resume_queue, job_description, results))
        thread.start()
        threads.append(thread)

    # Wait for all threads to complete
    for thread in threads:
        thread.join()

    # Update response_data with results
    for filename, result in results.items():
        response_data[filename]['key_feature'] = result['resume_key_aspect']
        response_data[filename]['score'] = result['resume_score']

    return response_data


@app.get("/download/{file_name}")
def download_file(file_name: str):
    """Endpoint to download a file by its name."""
    # file_path = os.path.join("extracted_files", file_name)  # Adjust the path as needed
    if os.path.exists(file_name):
        print(file_name)
        return FileResponse(file_name, media_type='application/octet-stream', filename = file_name.split('_', 2)[-1])
    else:
        return JSONResponse(content={"message": "File not found."}, status_code=404)
