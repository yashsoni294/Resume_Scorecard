from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse
from PyPDF2 import PdfReader
from docx import Document
import win32com.client
import os
import io
import zipfile
import pandas as pd
import time
import tkinter as tk
from tkinter import filedialog
from dotenv import load_dotenv
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from datetime import datetime

# Load API Key
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=API_KEY)


# Constants
TEMPLATES = {
    "job_description": """
        
        The text below is a job description:

        {job_description_text}

        Your task is to analyze the job description and extract key aspects that will help evaluate a candidate's suitability. Organize the extracted information into the following categories to support scoring. Also Avoid including unnecessary details or assumptions on your own :

        1. Candidate Profile :-
            Job-Related Keywords:
            Extract the most relevant keywords and phrases used in the job description, such as specific skills, technologies, or qualifications, that highlight the core expectations for the role.

            Relevant Past Roles and Responsibilities:
            Identify the types of roles and responsibilities that align closely with this position, based on the job description.

            Actionable Responsibilities:
            List clear, action-oriented tasks or expectations (e.g., "Develop and manage X," "Collaborate on Y") that indicate measurable outcomes for success in this role.

        2. Experience Section :-
            Years of Experience:
            Specify the minimum and preferred years of experience required or desired for the role.

            Technical Skills:
            Extract a comprehensive list of both mandatory and preferred technical skills, including tools, technologies, programming languages, or domain-specific expertise mentioned in the job description.

            Soft Skills and Teamwork:
            Identify soft skills and teamwork-related requirements (e.g., leadership, communication, collaboration), with examples if provided in the description.

        3. Educational Qualifications and Certifications :- 
            Minimum Educational Qualifications:
            Note the required or preferred educational qualifications for this role.

            Relevant Certifications/Training Programs:
            List certifications, licenses, or training programs that are directly or partially relevant to the job description.

            Output:
            Provide the extracted information in a concise, bullet-pointed format, ensuring relevance to the scoring criteria. Avoid including unnecessary details or assumptions on your own. This structured output will guide the scoring process.        
            """,

    "resume":""" 
            The text below is a resume:

            {resume_text}

            Your task is to extract information from the resume by focusing on the following areas. Ensure the extracted information is clear, structured, and concise also Ensure the output focuses only on the resume content without making assumptions or adding extra information from your own :

            ### 1. **Candidate Profile**
            - **Keywords Identified:** List relevant keywords found in the resume that reflect the candidate skills, roles, and expertise.  
            - **Past Roles Summary:** Provide an overview of the candidate's past roles and responsibilities, emphasizing their clarity and detail. Highlight the use of action words (e.g., "Developed," "Managed") and measurable outcomes.  
            - **Clarity of Responsibilities:** Note how clearly responsibilities are described, focusing on structured descriptions and measurable achievements.

            ### 2. **Experience Section**
            - **Years of Experience:** Indicate the total years of experience and specify industries or domains the candidate has worked in.  
            - **Technical Skills:** Identify technical skills explicitly mentioned in the resume, along with any supporting examples or certifications.  
            - **Communication and Teamwork Skills:** Highlight references to teamwork and communication skills, particularly examples such as leadership roles, collaboration efforts, or achievements in group settings.

            ### 3. **Educational Qualifications and Certifications**
            - **Educational Background:** Provide the highest qualification achieved and any notable academic achievements.  
            - **Certifications and Training:** List certifications or training programs mentioned in the resume, along with their relevance to enhancing the candidate expertise.

            Provide a structured summary of the extracted information in bullet points or short sentences for each section.

            """,
    "score": """
        Your task is to evaluate how well the provided resume aligns with the given job description by assessing three key factors: 1. Candidate Profile 2. Experience section 3. Educational Qualifications and Certifications. Based on this evaluation, provide a final score between 0 and 100, representing the overall suitability of the candidate for the job.

        ### Inputs:
        - **Resume Text:**  

        {resume_text}

        - **Job Description Text:**  

        {job_description}

        ### Scoring Guidelines:
        Give marks to the resume against the job description across the following criterias :-

        1. Candidate Profile (Max 15 Marks) :-

            •	Job-Related Keywords (Max 5 Marks):
                o	5 Points: IF Resume includes highly relevant keywords from the Job Description, indicating alignment with job requirements.
                o	3 Points: IF Resume Includes some relevant keywords but lacks critical ones from the Job Description.
                o	1 Point: IF Few or no relevant keywords used.
            •	Relevance of Past Roles to Job Description (Max 5 Marks):
                o	5 Points: IF Past roles and responsibilities align strongly with Job Description requirements.
                o	3 Points: IF Moderate alignment, some responsibilities match the Job Description.
                o	1 Point: IF Weak or no relevance to the Job Description.
            •	Clarity of Responsibilities (Max 5 Marks):
                o	5 Points: IF Responsibilities are clearly defined using action words (e.g., "Developed," "Managed") and measurable outcomes.
                o	3 Points: IF Responsibilities are described but lack action words or outcomes.
                o	1 Point: IF Responsibilities are vague or generic.
        
        2. Experience section (Max 65) :- 

            • Years of Experience (Max 15 Marks):
                o	15 Points: IF Experience matches or exceeds the Job Description requirements.
                o	10 Points: IF Slightly below required years but relevant experience.
                o	5 Points: IF Limited relevance or inadequate years of experience.
            • Matching Technical Skills (Max 40 Marks):
                o	40 Points: IF All technical skills mentioned in the Job Description are evident, with examples or certifications provided.
                o	30 Points: IF Most technical skills are evident but lack depth or examples.
                o	20 Points: IF Some technical skills match; others are missing.
                o	10 Points or Below: IF Minimal or no alignment with required technical skills.
            • Communication and Teamwork (Max 10 Marks):
                o	10 Points: IF Resume highlights soft skills with examples (e.g., "Led a team of 5," "Facilitated cross-functional communication").
                o	6 Points: IF Mentions soft skills but lacks examples.
                o	3 Points: IF Minimal mention of soft skills.

        3. Educational Qualifications and Certifications (Max 20) :- 
                
            • Meets Minimum Educational Qualifications (Max 15 Marks):
                o	15 Points: IF Meets or exceeds educational qualifications specified in the Job Description.
                o	10 Points: IF Meets basic qualifications but lacks advanced or desired qualifications.
                o	5 Points: IF Does not fully meet educational qualifications.
            • Additional Certifications/Training Programs (Max 5 Marks):
                o	5 Points: IF Certifications/training directly related to Job Description (e.g., HR certifications, technical tools training).
                o	3 Points: IF Certifications or training partially relevant to the Job Description.
                o	1 Point: IF No additional certifications.

        Add marks of all Three section (Candidate Profile, Experience section, Educational Qualifications and Certifications) Round the result to the nearest whole number.

        ### Output:
        Provide only the final average score as a single number (0-100) with no additional text or explanation.
        Never give any text as output give 0 score instead.
    """
    ,
}

app = FastAPI()

def get_conversation(template):
    """
    Initializes a conversation using a template and LangChain's LLM integration.
    Args:
        template (str): Template content for the conversation.
    Returns:
        Callable: Configured conversation object.
    """
    llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-flash",
        temperature=0.1,
        max_tokens=None,
        timeout=None,
        max_retries=3
    )
    prompt = PromptTemplate.from_template(template)
    return prompt | llm

def read_pdf(file: io.BytesIO):
    """Extract text from a PDF file."""
    pdf_reader = PdfReader(file)
    extracted_text = ""
    for page in pdf_reader.pages:
        extracted_text += page.extract_text()
    return extracted_text.strip()

def read_docx(file: io.BytesIO):
    """Extract text from a DOCX file."""
    document = Document(file)
    extracted_text = ""
    for paragraph in document.paragraphs:
        extracted_text += paragraph.text + "\n"
    return extracted_text.strip()

def read_doc(file_path: str):
    """
    Extract text from a DOC file using COM automation (Windows only).
    """
    try:
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(file_path))
        extracted_text = doc.Content.Text
        doc.Close(False)
        word.Quit()
        return extracted_text.strip()
    except Exception as e:
        return f"Error processing DOC file: {str(e)}"

def read_txt(file: io.BytesIO):
    """Extract text from a plain text file."""
    try:
        contents = file.read()
        return contents.decode("utf-8").strip()
    except Exception as e:
        return f"Error processing TXT file: {str(e)}"

def process_file(file_name: str, file_content: bytes):
    """
    Determine the file type and extract text based on its format.
    """
    try:
        file_extension = file_name.split(".")[-1].lower()
        file_stream = io.BytesIO(file_content)

        if file_extension == "pdf":
            return read_pdf(file_stream)
        elif file_extension == "docx":
            return read_docx(file_stream)
        elif file_extension == "doc":
            temp_file_path = f"temp_{file_name}"
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(file_content)
            text = read_doc(temp_file_path)
            os.remove(temp_file_path)
            return text
        elif file_extension == "txt":
            return read_txt(file_stream)
        
    except Exception as e:
        return f"Error processing file {file_name}: {str(e)}"


@app.post("/upload-files/")
async def upload_files(job_description: str, files: list[UploadFile] = File(...)):
    response_data = {}

    conversation_jd = get_conversation(TEMPLATES["job_description"])
    jd_response = conversation_jd.invoke({"job_description_text": job_description})
    processed_jd = jd_response.content
    print(processed_jd)

    for file in files:
        try:
            # Fallback to file extension if MIME type is not reliable
            file_extension = file.filename.split(".")[-1].lower()

            if file.content_type == "application/zip" or file_extension == "zip":
                zip_data = await file.read()
                zip_file = io.BytesIO(zip_data)
                extracted_files = []
                with zipfile.ZipFile(zip_file, 'r') as z:
                    # Create a directory for extracted files (optional)
                    extract_path = "extracted_files"
                    os.makedirs(extract_path, exist_ok=True)
                    
                    # Extract files
                    z.extractall(extract_path)
                    extracted_files = z.namelist()

                pdf_contents = {}
                for file_name in extracted_files:
                    file_path = os.path.join(extract_path, file_name)
                    if file_name.endswith(".pdf"):
                        file_path = os.path.join(extract_path, file_name)
                        with open(file_path, "rb") as pdf_file:
                            reader = PdfReader(pdf_file)
                            text = ""
                            for page in reader.pages:
                                text += page.extract_text()
                            # pdf_contents[file_name] = text
                            response_data[file_name] = {"content": text}
                    
                    # Process TXT files
                    elif file_name.endswith(".txt"):
                        with open(file_path, "r", encoding="utf-8") as txt_file:
                            text = txt_file.read()
                            response_data[file_name] = {"content": text}

                    elif file_name.endswith(".docx"):
                        try:
                            doc = Document(file_path)
                            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                            response_data[file_name] = {"content": text}
                        except Exception as e:
                            response_data[file_name] = {"content": e}

                    # Process DOC files
                    elif file_name.endswith(".doc"):
                        try:
                            word = win32com.client.Dispatch("Word.Application")
                            doc = word.Documents.Open(os.path.abspath(file_path))
                            text = doc.Content.Text
                            doc.Close()
                            word.Quit()
                            response_data[file_name] = {"content": text}  # Extracted text
                        except Exception as e:
                            response_data[file_name] = {"content": f"Error reading DOC file: {str(e)}"}


                # response_data[file.filename] = process_zip(file)
            elif file.content_type in ["application/x-rar-compressed", "application/vnd.rar"] or file_extension == "rar":
                # response_data[file.filename] = process_rar(file)
                pass
            elif file.content_type == "application/pdf" or file_extension == "pdf":
                extracted_text = read_pdf(io.BytesIO(file.file.read()))
                response_data[file.filename] = {"content": extracted_text}
            elif file.content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"] or file_extension in ["docx", "doc"]:
                if file_extension == "docx":
                    extracted_text = read_docx(io.BytesIO(file.file.read()))
                elif file_extension == "doc":
                    temp_file_path = f"temp_{file.filename}"
                    with open(temp_file_path, "wb") as temp_file:
                        temp_file.write(file.file.read())
                    extracted_text = read_doc(temp_file_path)
                    os.remove(temp_file_path)
                else:
                    raise ValueError("Unsupported Word file format.")
                response_data[file.filename] = {"content": extracted_text}
            elif file.content_type == "text/plain" or file_extension == "txt":
                extracted_text = read_txt(io.BytesIO(file.file.read()))
                response_data[file.filename] = {"content": extracted_text}
            else:
                response_data[file.filename] = {
                    "error": "Unsupported file type. Supported formats: .pdf, .docx, .doc, .txt, .zip, .rar"
                }

        except Exception as e:
            response_data[file.filename] = {
                "error": str(e)
            }
    conversation_resume = get_conversation(TEMPLATES["resume"])
    conversation_score = get_conversation(TEMPLATES["score"])
    for filename in response_data:
        print(f"{filename}") #: {response_data[filename]["content"]}

        resume_response = conversation_resume.invoke({"resume_text": response_data[filename]["content"]})
        time.sleep(3)
        # print(resume_response.content)
        response_data[filename]["key_feature"] = resume_response.content
        # response_data[filename] = {"key_aspect": resume_response.content}
        score_response = conversation_score.invoke({
            "resume_text": resume_response.content,
            "job_description": processed_jd
        })
        time.sleep(3)
        response_data[filename]["score"] = score_response.content
        print(score_response.content)

    return response_data

